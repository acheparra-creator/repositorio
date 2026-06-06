#!/usr/bin/env python3
"""
Reformateador de Expedientes Clínicos NOM-004-SSA3-2012
Uso: python3 procesar_expediente.py <ruta_expediente.docx>
- Texto negro = información original del médico
- Texto azul  = secciones agregadas para cumplir NOM-004
- Al inicio   = tabla comparativa de creatinina, TFG, proteinuria
"""
import sys, json, os, re, subprocess
from pathlib import Path

try:
    from docx import Document
    from docx.shared import RGBColor, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    os.system("pip3 install python-docx -q")
    from docx import Document
    from docx.shared import RGBColor, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

BLUE  = RGBColor(0x00, 0x70, 0xC0)
BLACK = RGBColor(0x00, 0x00, 0x00)
NAVY  = RGBColor(0x1F, 0x49, 0x7D)
GRAY  = RGBColor(0x99, 0x99, 0x99)

# ── Prompts ────────────────────────────────────────────────────────────────────
PROMPT_HISTORIA = """Eres experto en NOM-004-SSA3-2012 (expediente clínico mexicano).

TAREA: Reorganizar el expediente al formato NOM-004. Incluye historia clínica completa y las notas de evolución que estén en el texto.

REGLA:
- "added": false = información que SÍ estaba en el original
- "added": true  = información que NO estaba (usa "No registrado", "Pendiente", "[Inferido]: ...")
- NO inventes datos clínicos. SÍ agrega CIE-10, unidades faltantes, campos estructurales vacíos.

ESTRUCTURA:
h1: HISTORIA CLÍNICA
  h2: FICHA DE IDENTIFICACIÓN → fields: Nombre/FN/Edad/Sexo/Ocupación/Residencia/Derechohabiencia/Médico/Fecha
  h2: ANTECEDENTES HEREDO-FAMILIARES → text
  h2: ANTECEDENTES PERSONALES NO PATOLÓGICOS → fields: Tabaquismo/Alcoholismo/Alergias/Inmunizaciones/Hemotipo/AINES
  h2: ANTECEDENTES PERSONALES PATOLÓGICOS → text
  h2: PADECIMIENTO ACTUAL / MOTIVO DE CONSULTA → text
  h2: INTERROGATORIO POR APARATOS Y SISTEMAS → fields: Cardiovascular/Respiratorio/Digestivo/Urinario/Neurológico/Musculoesquelético
  h2: EXPLORACIÓN FÍSICA
    h3: Signos vitales → fields: TA/FC/FR/T/Peso/Talla/IMC
    h3: Exploración regional → text
  h2: ESTUDIOS DE LABORATORIO Y GABINETE → text
  h2: DIAGNÓSTICO → text con CIE-10
  h2: PLAN TERAPÉUTICO
    h3: Tratamiento farmacológico → medication por fármaco
    h3: Indicaciones no farmacológicas → text
  h2: PRONÓSTICO → text

Para cada consulta subsecuente en el texto (formato SOAP):
h1: NOTA DE EVOLUCIÓN — DD/MM/AAAA
  h2: S — Subjetivo → text
  h2: O — Objetivo
    h3: Signos vitales → fields TA/FC/FR/T/Peso
    h3: Estudios de laboratorio → text
    h3: Gabinete → text (solo si hay)
  h2: A — Evaluación → text con CIE-10
  h2: P — Plan
    h3: Tratamiento farmacológico → medication por fármaco
    h3: Indicaciones → text
    field: Próxima cita
divider ← entre cada nota de evolución

TABLA: extrae TODOS los valores de creatinina (mg/dL), TFG/depuración (ml/min), proteinuria (mg o g/día).

RESPUESTA: SOLO JSON válido, sin markdown:
{"blocks":[{"type":"h1","text":"...","added":false},{"type":"h2","text":"...","added":false},{"type":"h3","text":"...","added":false},{"type":"field","label":"...","value":"...","added":false},{"type":"text","text":"...","added":false},{"type":"medication","text":"...","added":false},{"type":"divider"}],"lab_table":[{"fecha":"MM/AAAA","creatinina":"valor","tfg":"valor","proteinuria":"valor"}]}"""

PROMPT_NOTAS = """Eres experto en NOM-004-SSA3-2012. Estas son notas de evolución de un expediente ya iniciado.

TAREA: Convierte cada nota al formato SOAP de NOM-004.

REGLA: "added":false = estaba en el original. "added":true = lo agregas (usa "No registrado", CIE-10, etc.)

Por cada consulta (formato SOAP):
h1: NOTA DE EVOLUCIÓN — DD/MM/AAAA
  h2: S — Subjetivo → text
  h2: O — Objetivo
    h3: Signos vitales → fields TA/FC/FR/T/Peso
    h3: Estudios de laboratorio → text
    h3: Gabinete → text (solo si hay)
  h2: A — Evaluación → text con CIE-10
  h2: P — Plan
    h3: Tratamiento farmacológico → medication por fármaco
    h3: Indicaciones → text
    field: Próxima cita
divider ← entre cada nota

TABLA: creatinina (mg/dL), TFG (ml/min), proteinuria de estas consultas solamente.

RESPUESTA: SOLO JSON:
{"blocks":[...],"lab_table":[...]}"""

MAX_CHARS = 7000  # chars per chunk


# ── Extracción ─────────────────────────────────────────────────────────────────
def extract_text(path):
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t: parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            r = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if r: parts.append(f"[TABLA] {r}")
    return "\n".join(parts)


# ── Chunking ───────────────────────────────────────────────────────────────────
def split_by_visits(text):
    """Split at date boundaries, keeping chunks under MAX_CHARS."""
    if len(text) <= MAX_CHARS:
        return [text]

    date_re = re.compile(r'(?m)^(\d{1,2}/\d{1,2}/\d{2,4})', re.MULTILINE)
    positions = [m.start() for m in date_re.finditer(text)]

    if not positions:
        return [text]

    chunks = []
    chunk_start = 0

    for pos in positions:
        if pos - chunk_start >= MAX_CHARS:
            chunks.append(text[chunk_start:pos])
            chunk_start = pos

    chunks.append(text[chunk_start:])
    return [c for c in chunks if c.strip()]


# ── Llamada a Claude ───────────────────────────────────────────────────────────
def claude_call(system_prompt, content, timeout=240):
    full_prompt = f"{system_prompt}\n\nTexto a procesar:\n\n{content}"
    result = subprocess.run(
        ["claude", "--model", "claude-sonnet-4-6", "-p", full_prompt],
        capture_output=True, text=True, timeout=timeout
    )
    if result.returncode != 0:
        raise RuntimeError(f"claude error: {result.stderr[:300]}")
    raw = result.stdout.strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```(?:json)?\n?", "", raw)
        raw = re.sub(r"\n?```$", "", raw.strip())
    return json.loads(raw)


def call_claude(text):
    """Process full expediente, splitting into chunks if needed."""
    chunks = split_by_visits(text)

    if len(chunks) == 1:
        print(f"    Procesando en 1 bloque (~{len(text)} chars)...")
        return claude_call(PROMPT_HISTORIA, text)

    print(f"    Expediente grande: procesando en {len(chunks)} bloques...")
    
    # First chunk: full historia clínica structure
    print(f"    Bloque 1/{len(chunks)} — historia clínica + primeras consultas...")
    data = claude_call(PROMPT_HISTORIA, chunks[0])

    # Subsequent chunks: evolution notes only
    for i, chunk in enumerate(chunks[1:], 2):
        print(f"    Bloque {i}/{len(chunks)} — notas de evolución...")
        chunk_data = claude_call(PROMPT_NOTAS, chunk)
        if chunk_data.get("blocks"):
            data["blocks"].append({"type": "divider"})
            data["blocks"].extend(chunk_data["blocks"])
        data["lab_table"].extend(chunk_data.get("lab_table", []))

    return data


# ── Helpers de documento ───────────────────────────────────────────────────────
def col(added):
    return BLUE if added else BLACK


def build_lab_table(doc, lab_table):
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("TABLA COMPARATIVA DE LABORATORIALES RENALES")
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s.add_run("Creatinina sérica  ·  Tasa de Filtrado Glomerular  ·  Proteinuria")
    r2.font.size = Pt(10); r2.font.color.rgb = NAVY

    if not lab_table:
        p = doc.add_paragraph()
        ri = p.add_run("No se encontraron registros de laboratorio en el expediente.")
        ri.italic = True
        doc.add_paragraph(); return

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Fecha", "Creatinina (mg/dL)", "TFG (ml/min)", "Proteinuria"]):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True; r.font.color.rgb = NAVY
    for e in lab_table:
        row = tbl.add_row().cells
        row[0].text = e.get("fecha", "N/D")
        row[1].text = e.get("creatinina", "N/D")
        row[2].text = e.get("tfg", "N/D")
        row[3].text = e.get("proteinuria", "N/D")
    doc.add_paragraph()


def build_document(data, output_path):
    doc = Document()
    sec = doc.sections[0]
    for attr in ("left_margin","right_margin","top_margin","bottom_margin"):
        setattr(sec, attr, Cm(2.5))
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    build_lab_table(doc, data.get("lab_table", []))
    doc.add_page_break()

    added_count = 0
    for block in data.get("blocks", []):
        btype = block.get("type", "text")
        added = block.get("added", False)
        if added: added_count += 1

        if btype == "h1":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(block.get("text", ""))
            r.bold = True; r.font.size = Pt(14)
            r.font.color.rgb = BLUE if added else NAVY

        elif btype == "h2":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(block.get("text", ""))
            r.bold = True; r.font.size = Pt(12)
            r.font.color.rgb = BLUE if added else NAVY

        elif btype == "h3":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            r = p.add_run(block.get("text", ""))
            r.bold = True; r.font.size = Pt(11)
            r.font.color.rgb = col(added)

        elif btype == "field":
            p = doc.add_paragraph()
            lr = p.add_run(f"{block.get('label','')}: ")
            lr.bold = True; lr.font.color.rgb = col(added)
            vr = p.add_run(block.get("value", ""))
            vr.font.color.rgb = col(added)

        elif btype == "text":
            p = doc.add_paragraph()
            r = p.add_run(block.get("text", ""))
            r.font.color.rgb = col(added)

        elif btype == "medication":
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(block.get("text", ""))
            r.font.color.rgb = col(added)

        elif btype == "divider":
            p = doc.add_paragraph("─" * 70)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            for r in p.runs:
                r.font.color.rgb = GRAY

    doc.save(output_path)
    return added_count


def main():
    if len(sys.argv) < 2:
        print("Uso: python3 procesar_expediente.py <ruta_expediente.docx>")
        sys.exit(1)
    path = sys.argv[1]
    if not os.path.exists(path):
        print(f"❌  Archivo no encontrado: {path}"); sys.exit(1)
    if not path.lower().endswith(".docx"):
        print("❌  El archivo debe ser .docx"); sys.exit(1)

    print(f"\n📋  Expediente: {Path(path).stem}")
    print("    Leyendo contenido...")
    text = extract_text(path)
    print(f"    Tamaño: {len(text)} caracteres")

    try:
        data = call_claude(text)
    except json.JSONDecodeError as e:
        print(f"❌  Error en JSON de respuesta: {e}"); sys.exit(1)
    except RuntimeError as e:
        print(f"❌  {e}"); sys.exit(1)

    lab_count = len(data.get("lab_table", []))
    print("    Generando documento reformateado...")
    added_count = build_document(data, path)

    print(f"\n✅  ¡Listo! → {path}")
    print(f"    📊  Registros en tabla laboratoriales: {lab_count}")
    print(f"    🔵  Elementos en azul (NOM-004 agregados): {added_count}")
    print(f"    ⚫  Texto negro = información original\n")

if __name__ == "__main__":
    main()
